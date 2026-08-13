#!/usr/bin/env python3
"""Génère le guide Word OSINT — auteur ABOU SAYABOU."""

from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

OUT = Path(__file__).resolve().parent / "Guide_OSINT_10_Outils_ABOU_SAYABOU.docx"

NAVY = RGBColor(0x0D, 0x3B, 0x66)
BLUE = RGBColor(0x1F, 0x4E, 0x79)
ACCENT = RGBColor(0x2E, 0x75, 0xB6)
DARK = RGBColor(0x2B, 0x2B, 0x2B)
GRAY = RGBColor(0x55, 0x55, 0x55)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
GREEN = RGBColor(0x1B, 0x5E, 0x20)


def set_run_font(run, name="Calibri", size=11, bold=False, italic=False, color=DARK):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    run.font.color.rgb = color


def shade_cell(cell, hex_color: str):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), hex_color)
    shd.set(qn("w:val"), "clear")
    tcPr.append(shd)


def set_cell_border(cell, **kwargs):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    for edge in ("top", "left", "bottom", "right"):
        element = OxmlElement(f"w:{edge}")
        element.set(qn("w:val"), kwargs.get(edge, "single"))
        element.set(qn("w:sz"), kwargs.get("sz", "4"))
        element.set(qn("w:color"), kwargs.get("color", "1F4E79"))
        tcBorders.append(element)
    tcPr.append(tcBorders)


def set_header_footer(doc: Document):
    section = doc.sections[0]
    header = section.header
    header.is_linked_to_previous = False
    p = header.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = p.add_run("Guide OSINT — 10 outils incontournables  |  ABOU SAYABOU")
    set_run_font(run, size=9, italic=True, color=GRAY)

    footer = section.footer
    footer.is_linked_to_previous = False
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Document interne — usage pédagogique et professionnel  •  Confidentiel  •  Page ")
    set_run_font(run, size=8, color=GRAY)

    # PAGE field
    fldChar1 = OxmlElement("w:fldChar")
    fldChar1.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    fldChar2 = OxmlElement("w:fldChar")
    fldChar2.set(qn("w:fldCharType"), "end")
    r = p.add_run()
    r._r.append(fldChar1)
    r._r.append(instr)
    r._r.append(fldChar2)
    set_run_font(r, size=8, color=GRAY)


def add_heading_styled(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = NAVY if level == 1 else BLUE
        run.font.name = "Calibri"
    return h


def add_para(doc, text, *, size=11, bold=False, italic=False, color=DARK, align="left", space_after=8):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.line_spacing = 1.15
    if align == "center":
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    elif align == "justify":
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    else:
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(text)
    set_run_font(run, size=size, bold=bold, italic=italic, color=color)
    return p


def add_bullet(doc, text, bold_prefix=None):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.left_indent = Cm(1)
    if bold_prefix:
        r = p.add_run(bold_prefix)
        set_run_font(r, size=11, bold=True, color=BLUE)
        r2 = p.add_run(text)
        set_run_font(r2, size=11, color=DARK)
    else:
        r = p.add_run(text)
        set_run_font(r, size=11, color=DARK)
    return p


def add_code(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.left_indent = Cm(0.4)
    p.paragraph_format.line_spacing = 1.08
    run = p.add_run(text)
    set_run_font(run, name="Consolas", size=9, color=RGBColor(0x1A, 0x1A, 0x1A))
    # light background via shading on paragraph
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), "F4F7FA")
    shd.set(qn("w:val"), "clear")
    pPr.append(shd)
    return p


def add_callout(doc, title, body):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.cell(0, 0)
    shade_cell(cell, "E8F0F8")
    set_cell_border(cell, color="1F4E79")
    p = cell.paragraphs[0]
    r = p.add_run(title + "  ")
    set_run_font(r, size=10, bold=True, color=NAVY)
    r2 = p.add_run(body)
    set_run_font(r2, size=10, color=DARK)
    doc.add_paragraph()


def add_kv_table(doc, rows, headers=("Élément", "Détail")):
    table = doc.add_table(rows=1, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = ""
        p = hdr[i].paragraphs[0]
        r = p.add_run(h)
        set_run_font(r, size=10, bold=True, color=WHITE)
        shade_cell(hdr[i], "1F4E79")
        set_cell_border(hdr[i])
    for i, (k, v) in enumerate(rows):
        row = table.add_row().cells
        row[0].text = ""
        p = row[0].paragraphs[0]
        r = p.add_run(k)
        set_run_font(r, size=10, bold=True, color=NAVY)
        row[1].text = ""
        p = row[1].paragraphs[0]
        r = p.add_run(v)
        set_run_font(r, size=10, color=DARK)
        bg = "F7FAFC" if i % 2 == 0 else "FFFFFF"
        shade_cell(row[0], bg)
        shade_cell(row[1], bg)
        set_cell_border(row[0], color="BDD3E6")
        set_cell_border(row[1], color="BDD3E6")
    doc.add_paragraph()
    return table


def add_page_break(doc):
    doc.add_page_break()


TOOLS = [
    {
        "num": "01",
        "name": "OSINT Framework",
        "type": "Annuaire / cartographie de ressources",
        "licence": "Site public (ressources majoritairement gratuites)",
        "url": "https://osintframework.com",
        "role": (
            "Point d’entrée de toute investigation. OSINT Framework n’est pas un scanner : "
            "c’est une carte mentale interactive qui classe des centaines de sites et d’outils "
            "selon le type de donnée recherchée (e-mail, nom d’utilisateur, domaine, image, "
            "géolocalisation, dark web, etc.). Un expert s’en sert pour ne rien oublier et pour "
            "choisir l’outil adapté au cas d’usage."
        ),
        "install": [
            "Aucune installation : accès via navigateur.",
            "Miroir / projet : https://github.com/lockfale/OSINT-Framework",
            "Option : cloner le dépôt pour usage hors ligne (HTML/JS statique).",
        ],
        "install_cmd": (
            "git clone https://github.com/lockfale/OSINT-Framework.git\n"
            "cd OSINT-Framework\n"
            "# Ouvrir index.html dans un navigateur"
        ),
        "usage": [
            "Ouvrir osintframework.com.",
            "Choisir une branche selon l’objectif (Username, Email Address, Domain Name, IP Address…).",
            "Déplier les nœuds jusqu’à l’outil, puis ouvrir le lien.",
            "Noter les sources utilisées (traçabilité de l’enquête).",
        ],
        "cas": [
            "Démarrage d’un dossier d’investigation (personne, organisation, incident).",
            "Recherche d’alternatives quand un outil principal est indisponible.",
            "Formation / check-list méthodologique d’une équipe SOC ou CERT.",
        ],
        "limites": (
            "Les liens évoluent ; certains outils deviennent payants ou disparaissent. "
            "Toujours vérifier la légalité et les CGU avant collecte."
        ),
    },
    {
        "num": "02",
        "name": "SpiderFoot",
        "type": "Plateforme d’automatisation OSINT",
        "licence": "MIT (open source)",
        "url": "https://github.com/smicallef/spiderfoot",
        "role": (
            "SpiderFoot automatise la collecte et la corrélation d’informations sur une cible "
            "(domaine, IP, e-mail, nom d’utilisateur, ASN, Bitcoin, etc.) via plus de 200 modules. "
            "C’est l’outil de « premier balayage » pour cartographier une surface d’attaque "
            "ou préparer une analyse de menace."
        ),
        "install": [
            "Prérequis : Python 3.7+ (idéalement 3.10+), Git.",
            "Installation recommandée : clone GitHub + pip.",
            "Interface web locale après lancement du serveur.",
        ],
        "install_cmd": (
            "git clone https://github.com/smicallef/spiderfoot.git\n"
            "cd spiderfoot\n"
            "pip3 install -r requirements.txt\n"
            "python3 ./sf.py -l 127.0.0.1:5001"
        ),
        "usage": [
            "Ouvrir http://127.0.0.1:5001 dans le navigateur.",
            "New Scan → nom du scan + cible (ex. exemple.com).",
            "Choisir un use-case (Footprint, Investigate, Passive) ou des modules précis.",
            "Lancer le scan, puis explorer Graph / Correlations / Browse.",
            "Exporter (CSV, JSON, GEXF) pour Maltego ou un rapport.",
        ],
        "cas": [
            "Reconnaissance passive d’un SI avant pentest (périmètre autorisé).",
            "Enrichissement d’IOC (IP, domaine) en threat intelligence.",
            "Détection d’e-mails, fuites, mentions, technologies exposées.",
        ],
        "limites": (
            "Beaucoup de modules sont plus riches avec des clés API (Shodan, VirusTotal, HIBP…). "
            "Limiter le scan (use-case Passive) pour rester dans un cadre légal et discret."
        ),
    },
    {
        "num": "03",
        "name": "theHarvester",
        "type": "Outil de reconnaissance (e-mails, sous-domaines, hosts)",
        "licence": "GNU GPL (open source)",
        "url": "https://github.com/laramies/theHarvester",
        "role": (
            "theHarvester interroge des sources publiques (moteurs, certificats, API) pour extraire "
            "des e-mails, sous-domaines, noms d’hôtes et adresses IP liés à un domaine. "
            "C’est un classique Kali Linux, rapide à lancer en ligne de commande."
        ),
        "install": [
            "Souvent préinstallé sur Kali Linux.",
            "Sinon : pipx ou clone GitHub.",
            "Certaines sources nécessitent des clés API (dans api-keys.yaml).",
        ],
        "install_cmd": (
            "# Kali / Debian\n"
            "sudo apt update && sudo apt install theharvester\n\n"
            "# Ou depuis les sources\n"
            "git clone https://github.com/laramies/theHarvester.git\n"
            "cd theHarvester\n"
            "python3 -m pip install -r requirements.txt\n"
            "python3 theHarvester.py -h"
        ),
        "usage": [
            "Syntaxe de base : theHarvester -d domaine -b source -l limite.",
            "Sources utiles : crtsh, duckduckgo, bing, hackertarget, rapiddns, urlscan.",
            "Ajouter -f rapport pour exporter XML/JSON/HTML.",
            "Combiner plusieurs sources : -b all (plus bruyant / plus lent).",
        ],
        "cas": [
            "Cartographier les e-mails publics d’une organisation (phishing / awareness).",
            "Découvrir des sous-domaines oubliés (shadow IT).",
            "Alimenter un inventaire d’assets avant un audit.",
        ],
        "limites": (
            "Les résultats dépendent des index publics ; ils ne sont jamais exhaustifs. "
            "Respecter le périmètre contractuel et éviter le scraping agressif."
        ),
        "exemples": (
            "theHarvester -d exemple.com -b crtsh,hackertarget -l 200\n"
            "theHarvester -d exemple.com -b bing -f rapport_exemple"
        ),
    },
    {
        "num": "04",
        "name": "Recon-ng",
        "type": "Framework OSINT modulaire (style Metasploit)",
        "licence": "GPL (open source)",
        "url": "https://github.com/lanmaster53/recon-ng",
        "role": (
            "Recon-ng fournit un shell interactif, des workspaces, une base SQLite et des modules "
            "pour enchaîner des recherches (WHOIS, DNS, réseaux sociaux, API). "
            "Il convient aux enquêtes reproductibles et à la constitution d’un dossier structuré."
        ),
        "install": [
            "Prérequis : Python 3.",
            "Installation via pip ou clone du dépôt officiel.",
            "Les modules marketplace se téléchargent ensuite dans le shell.",
        ],
        "install_cmd": (
            "pip3 install recon-ng\n"
            "# ou\n"
            "git clone https://github.com/lanmaster53/recon-ng.git\n"
            "cd recon-ng\n"
            "pip3 install -r REQUIREMENTS\n"
            "./recon-ng"
        ),
        "usage": [
            "Lancer recon-ng puis créer un workspace : workspaces create dossier1.",
            "Installer des modules : marketplace search ; marketplace install <module>.",
            "Charger un module : modules load recon/domains-hosts/…",
            "Définir les options : options set SOURCE exemple.com ; run.",
            "Consulter les résultats : show hosts ; show contacts ; db query SELECT * FROM hosts.",
            "Ajouter des clés API : keys add shodan_api <cle>.",
        ],
        "cas": [
            "Enquête structurée multi-sources avec historisation SQLite.",
            "Scripts / reporting pour une équipe CTI.",
            "Enrichissement progressif (domaine → hosts → contacts).",
        ],
        "limites": (
            "Courbe d’apprentissage (commandes type Metasploit). "
            "Sans clés API, de nombreux modules restent limités."
        ),
    },
    {
        "num": "05",
        "name": "OWASP Amass",
        "type": "Énumération de surface d’attaque / sous-domaines",
        "licence": "Apache 2.0 (open source) — projet OWASP",
        "url": "https://github.com/owasp-amass/amass",
        "role": (
            "Amass est la référence pour découvrir des sous-domaines, relations DNS, ASN et "
            "infrastructures associées. Il combine sources passives (certificats, archives) "
            "et techniques actives (si autorisées). Indispensable en pentest et ASM "
            "(Attack Surface Management)."
        ),
        "install": [
            "Binaire Go, paquet Kali, ou compilation depuis les sources.",
            "Fichier de configuration YAML pour les API (optionnel mais recommandé).",
        ],
        "install_cmd": (
            "# Kali\n"
            "sudo apt install amass\n\n"
            "# Go\n"
            "go install -v github.com/owasp-amass/amass/v4/...@master\n\n"
            "# Vérification\n"
            "amass -version"
        ),
        "usage": [
            "Mode enum : découverte de noms associés à un domaine.",
            "Mode intel : infos ASN, réseaux, organisation.",
            "Mode viz : graphe (D3, Maltego, Graphviz).",
            "Toujours commencer en passif, puis actif seulement avec autorisation.",
        ],
        "cas": [
            "Inventaire de sous-domaines (y compris takeover potentiels).",
            "Cartographie DNS / ASN d’une organisation.",
            "Préparation d’un pentest externe.",
        ],
        "limites": (
            "Les scans actifs (brute DNS, zone walking) sont détectables. "
            "Configurer rate-limit et rester dans le cadre légal / contrat."
        ),
        "exemples": (
            "amass enum -passive -d exemple.com -o amass_passif.txt\n"
            "amass intel -org \"Exemple SA\"\n"
            "amass viz -d3 -o graphe.html"
        ),
    },
    {
        "num": "06",
        "name": "Sherlock",
        "type": "Recherche de pseudos sur les réseaux sociaux",
        "licence": "MIT (open source)",
        "url": "https://github.com/sherlock-project/sherlock",
        "role": (
            "Sherlock teste un identifiant (username) sur des centaines de sites "
            "(GitHub, Twitter/X, Instagram, Reddit, forums, etc.) et indique où le compte existe. "
            "C’est un outil clé d’investigation d’identité numérique (SOCMINT)."
        ),
        "install": [
            "Python 3.8+ recommandé.",
            "pipx (isolé) ou clone GitHub.",
        ],
        "install_cmd": (
            "pipx install sherlock-project\n"
            "# ou\n"
            "git clone https://github.com/sherlock-project/sherlock.git\n"
            "cd sherlock\n"
            "python3 -m pip install -r requirements.txt"
        ),
        "usage": [
            "Commande : sherlock <username>",
            "Plusieurs identifiants : sherlock alice bob charlie",
            "Export : --csv ou --xlsx pour le dossier d’enquête.",
            "Filtrer sites : --site GitHub --site Reddit",
            "Timeout / proxy si besoin de discrétion réseau.",
        ],
        "cas": [
            "Corréler un pseudo d’attaquant / fraudeur sur plusieurs plateformes.",
            "Vérifier l’empreinte numérique d’une personne (cadre légal RH / enquête).",
            "Préparer une analyse SOCMINT avant interview ou IR.",
        ],
        "limites": (
            "Faux positifs possibles (pseudos réutilisés). "
            "Ne pas harceler, ne pas contourner l’authentification : uniquement l’existence publique de profils."
        ),
        "exemples": (
            "sherlock johndoe --csv\n"
            "sherlock johndoe --site GitHub --site Reddit --timeout 10"
        ),
    },
    {
        "num": "07",
        "name": "ExifTool",
        "type": "Analyse et édition de métadonnées de fichiers",
        "licence": "Open source (Phil Harvey)",
        "url": "https://exiftool.org",
        "role": (
            "ExifTool lit (et peut modifier) les métadonnées EXIF, IPTC, XMP, GPS, auteurs, "
            "logiciels, timestamps de photos, PDF, documents Office, etc. "
            "Indispensable en forensics, OSINT image et contrôle de fuites d’information."
        ),
        "install": [
            "Paquets Linux (libimage-exiftool-perl), Windows installer, macOS Homebrew.",
            "Aucun serveur : binaire en ligne de commande.",
        ],
        "install_cmd": (
            "# Debian / Ubuntu / Kali\n"
            "sudo apt install libimage-exiftool-perl\n\n"
            "# macOS\n"
            "brew install exiftool\n\n"
            "# Vérification\n"
            "exiftool -ver"
        ),
        "usage": [
            "Lecture complète : exiftool photo.jpg",
            "GPS uniquement : exiftool -gpslatitude -gpslongitude -n photo.jpg",
            "Récursif : exiftool -r dossier/",
            "Export CSV : exiftool -csv -r dossier/ > meta.csv",
            "Nettoyage (attention, destructif) : exiftool -all= fichier.pdf",
        ],
        "cas": [
            "Géolocaliser une photo publiée (si GPS non retiré).",
            "Identifier l’appareil / logiciel ayant produit un document.",
            "Vérifier qu’un livrable ne fuit pas de métadonnées internes (auteur, chemin UNC).",
        ],
        "limites": (
            "Les réseaux sociaux stripent souvent l’EXIF. "
            "Les métadonnées peuvent être falsifiées : les croiser avec d’autres preuves."
        ),
        "exemples": (
            "exiftool image.jpg\n"
            "exiftool -gpslatitude -gpslongitude -CreateDate -Make -Model image.jpg\n"
            "exiftool -csv -r ./preuves > metadonnees.csv"
        ),
    },
    {
        "num": "08",
        "name": "Maltego CE (Community Edition)",
        "type": "Analyse relationnelle et visualisation de graphes",
        "licence": "Community Edition gratuite (transforms limitées) ; versions payantes Pro/Organization",
        "url": "https://www.maltego.com",
        "role": (
            "Maltego transforme des entités (personne, domaine, IP, e-mail, AS) en graphe "
            "via des « transforms » (intégrations). C’est l’outil de référence pour voir "
            "les liens cachés et présenter une enquête à un décideur. "
            "La CE est gratuite avec inscription ; ce n’est pas du 100 % open source, "
            "mais c’est le standard que tout expert doit connaître."
        ),
        "install": [
            "Créer un compte Maltego Community.",
            "Télécharger le client (Windows / Linux / macOS) depuis maltego.com.",
            "Java / runtime fourni avec l’installateur moderne.",
            "Kali : paquet maltego souvent disponible.",
        ],
        "install_cmd": (
            "# Kali\n"
            "sudo apt update && sudo apt install maltego\n\n"
            "# Sinon : installer depuis https://www.maltego.com/downloads/\n"
            "# Puis lancer Maltego et se connecter avec le compte Community"
        ),
        "usage": [
            "Nouveau graphe → glisser une entité (Domain, Person, Email…).",
            "Clic droit → Run Transform (ex. To DNS Name, To Email Address).",
            "Étendre progressivement ; éviter de tout lancer d’un coup (bruit).",
            "Layout : Organic / Hierarchical pour la lisibilité.",
            "Exporter PNG/PDF pour le rapport d’investigation.",
            "Installer des hubs (transforms) supplémentaires si besoin (VirusTotal, etc.).",
        ],
        "cas": [
            "Lier infrastructure phishing (domaines, MX, IP, certificats).",
            "Cartographier une campagne (e-mails, alias, sites).",
            "Support visuel en threat intel / enquête fraude.",
        ],
        "limites": (
            "Quota de transforms en CE. Données parfois incomplètes. "
            "Ne pas confondre corrélation graphique et preuve juridique."
        ),
    },
    {
        "num": "09",
        "name": "crt.sh",
        "type": "Moteur de recherche Certificate Transparency",
        "licence": "Service public (Sectigo / crt.sh) — données CT ouvertes",
        "url": "https://crt.sh",
        "role": (
            "crt.sh indexe les certificats TLS publiés dans les journaux Certificate Transparency. "
            "On y trouve des sous-domaines, SAN, dates d’émission, émetteurs — même pour des "
            "hôtes jamais publiés sur le site web. Complément naturel d’Amass et theHarvester."
        ),
        "install": [
            "Aucune installation : interface web.",
            "Requêtes SQL possibles (PostgreSQL public) pour les profils avancés.",
            "API / URL de recherche : https://crt.sh/?q=%.exemple.com&output=json",
        ],
        "install_cmd": (
            "# Exemple curl (JSON)\n"
            "curl -s \"https://crt.sh/?q=%25.exemple.com&output=json\" | head\n\n"
            "# Filtrer des FQDN uniques (jq)\n"
            "curl -s \"https://crt.sh/?q=%25.exemple.com&output=json\" \\\n"
            "  | jq -r '.[].name_value' | sed 's/\\n/\\n/g' | sort -u"
        ),
        "usage": [
            "Saisir %.exemple.com ou exemple.com dans la barre de recherche.",
            "Examiner Identity / Logged At / Issuer.",
            "Repérer des environnements (staging, vpn, owa, dev) non documentés.",
            "Croiser avec Amass / DNS pour valider les hôtes encore actifs.",
        ],
        "cas": [
            "Découverte de sous-domaines via certificats.",
            "Suivi de l’émission de certificats (typosquatting, phishing).",
            "Audit de périmètre SSL d’une organisation.",
        ],
        "limites": (
            "Un certificat ne signifie pas que le service est encore en ligne. "
            "Wildcards (*.domaine) masquent une partie des noms."
        ),
    },
    {
        "num": "10",
        "name": "Wayback Machine (Internet Archive)",
        "type": "Archives historiques du Web",
        "licence": "Service public à but non lucratif (Internet Archive)",
        "url": "https://web.archive.org",
        "role": (
            "La Wayback Machine conserve des copies datées de pages web. "
            "On y retrouve des contenus retirés, d’anciennes pages « équipe / contact », "
            "des fichiers JS exposés, des mentions de partenaires ou des fuites passées. "
            "Outil majeur d’OSINT historique et de réponse à incident."
        ),
        "install": [
            "Aucune installation pour l’usage web.",
            "Outils complémentaires OSS : waybackurls, gau, waymore (pour extraire les URL archivées).",
        ],
        "install_cmd": (
            "# Extraction d’URL archivées (Go)\n"
            "go install github.com/tomnomnom/waybackurls@latest\n"
            "echo exemple.com | waybackurls > urls_archivees.txt\n\n"
            "# Alternative\n"
            "pipx install waymore\n"
            "waymore -i exemple.com -mode U -o waymore_out"
        ),
        "usage": [
            "Aller sur web.archive.org et coller l’URL.",
            "Choisir une date sur le calendrier (captures *).",
            "Comparer plusieurs millésimes (diff humain).",
            "Chercher /robots.txt, /sitemap.xml, anciennes pages login.",
            "Pour un domaine entier : CDX API ou waybackurls.",
        ],
        "cas": [
            "Reconstituer un site de phishing déjà hors ligne.",
            "Trouver d’anciens e-mails / organigrammes publics.",
            "Identifier des endpoints API ou JS disparus du site actuel.",
        ],
        "limites": (
            "Toutes les pages ne sont pas archivées (robots.txt, login, CAPTCHA). "
            "Le contenu peut être incomplet ou en cache partiel."
        ),
        "exemples": (
            "https://web.archive.org/web/*/https://exemple.com/*\n"
            "echo exemple.com | waybackurls | grep -Ei 'admin|backup|config'"
        ),
    },
]


def build():
    doc = Document()

    # Page setup
    section = doc.sections[0]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.left_margin = Cm(2.2)
    section.right_margin = Cm(2.2)
    set_header_footer(doc)

    # Styles Normal
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)
    style.font.color.rgb = DARK

    # ========== PAGE DE GARDE ==========
    for _ in range(3):
        doc.add_paragraph()

    banner = doc.add_table(rows=1, cols=1)
    cell = banner.cell(0, 0)
    shade_cell(cell, "0D3B66")
    set_cell_border(cell, color="0D3B66")
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("\nCYBERSÉCURITÉ  •  OPEN SOURCE INTELLIGENCE\n")
    set_run_font(r, size=12, bold=True, color=WHITE)
    p2 = cell.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p2.add_run("Les 10 outils et sites OSINT\nqu’un expert doit connaître")
    set_run_font(r, size=26, bold=True, color=WHITE)
    p3 = cell.add_paragraph()
    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p3.add_run("\nInstallation, usage opérationnel et cas d’emploi\n")
    set_run_font(r, size=13, italic=True, color=RGBColor(0xD6, 0xE3, 0xF0))

    doc.add_paragraph()
    add_para(doc, "DOCUMENT TECHNIQUE", size=12, bold=True, color=ACCENT, align="center", space_after=4)
    add_para(doc, "Auteur", size=11, color=GRAY, align="center", space_after=0)
    add_para(doc, "ABOU SAYABOU", size=22, bold=True, color=NAVY, align="center", space_after=6)
    add_para(
        doc,
        "Consultant / Expert en cybersécurité  •  OSINT, SOC, durcissement",
        size=11,
        italic=True,
        color=GRAY,
        align="center",
        space_after=16,
    )

    meta = doc.add_table(rows=4, cols=2)
    meta.alignment = WD_TABLE_ALIGNMENT.CENTER
    infos = [
        ("Version", "1.0"),
        ("Date", date.today().strftime("%d/%m/%Y")),
        ("Classification", "Interne — usage professionnel et pédagogique"),
        ("Périmètre", "10 outils OSINT open source / publics incontournables"),
    ]
    for i, (k, v) in enumerate(infos):
        c0, c1 = meta.rows[i].cells
        c0.text = ""
        r = c0.paragraphs[0].add_run(k)
        set_run_font(r, size=10, bold=True, color=WHITE)
        shade_cell(c0, "1F4E79")
        c1.text = ""
        r = c1.paragraphs[0].add_run(v)
        set_run_font(r, size=10, color=DARK)
        shade_cell(c1, "F4F7FA")
        set_cell_border(c0)
        set_cell_border(c1)

    add_para(
        doc,
        "\nCe document décrit, pour chaque outil : le rôle dans une enquête, "
        "l’installation, l’usage concret, des exemples de commandes, les cas d’emploi "
        "cyber et les limites / précautions légales.",
        size=10,
        italic=True,
        color=GRAY,
        align="justify",
    )

    add_page_break(doc)

    # ========== SOMMAIRE ==========
    add_heading_styled(doc, "Sommaire", 1)
    toc_items = [
        "1.  Objet, cadre légal et méthode",
        "2.  Vue d’ensemble des 10 outils",
        "3.  Fiches détaillées",
    ]
    for t in toc_items:
        add_para(doc, t, size=12, color=BLUE, space_after=4)
    for t in TOOLS:
        add_para(doc, f"      3.{int(t['num'])}  {t['name']}", size=11, color=DARK, space_after=2)
    add_para(doc, "4.  Stack minimale recommandée", size=12, color=BLUE, space_after=4)
    add_para(doc, "5.  Outils complémentaires", size=12, color=BLUE, space_after=4)
    add_para(doc, "6.  Bonnes pratiques d’expert", size=12, color=BLUE, space_after=4)
    add_para(doc, "7.  Références", size=12, color=BLUE, space_after=12)

    # ========== 1 ==========
    add_heading_styled(doc, "1. Objet, cadre légal et méthode", 1)
    add_heading_styled(doc, "1.1 Objet", 2)
    add_para(
        doc,
        "L’Open Source Intelligence (OSINT) consiste à collecter, vérifier et corréler des "
        "informations publiquement accessibles pour soutenir la cybersécurité : reconnaissance "
        "autorisée, threat intelligence, investigation d’incident, lutte anti-fraude, "
        "cartographie de surface d’attaque. Ce guide cible les dix ressources qu’un expert "
        "doit savoir installer et opérer, pas seulement citer.",
        align="justify",
    )

    add_heading_styled(doc, "1.2 Cadre d’emploi", 2)
    add_bullet(doc, "Utiliser uniquement des sources publiques ou des API autorisées.")
    add_bullet(doc, "Disposer d’un mandat (audit, pentest, IR, enquête interne) avant toute collecte ciblée.")
    add_bullet(doc, "Ne pas contourner authentification, CAPTCHA, ou mesures d’accès.")
    add_bullet(doc, "Tracer les sources (URL, date, outil, capture) pour la traçabilité.")
    add_bullet(doc, "Protéger les données personnelles (RGPD) : minimisation, finalité, durée de conservation.")
    add_callout(
        doc,
        "Avertissement.",
        "Les commandes d’exemple utilisent le domaine fictif exemple.com. "
        "Ne jamais lancer d’énumération active sur une cible réelle sans autorisation écrite.",
    )

    add_heading_styled(doc, "1.3 Méthode d’enquête (rappel)", 2)
    add_bullet(doc, "Cadrer le besoin (identité, infra, fuite, phishing).", bold_prefix="1. ")
    add_bullet(doc, "Choisir les sources via OSINT Framework.", bold_prefix="2. ")
    add_bullet(doc, "Collecter (passif d’abord) : theHarvester, Amass, crt.sh, Wayback.", bold_prefix="3. ")
    add_bullet(doc, "Automatiser / corréler : SpiderFoot, Recon-ng.", bold_prefix="4. ")
    add_bullet(doc, "Identité / artefacts : Sherlock, ExifTool.", bold_prefix="5. ")
    add_bullet(doc, "Visualiser et rapporter : Maltego CE.", bold_prefix="6. ")

    # ========== 2 ==========
    add_heading_styled(doc, "2. Vue d’ensemble des 10 outils", 1)
    table = doc.add_table(rows=1, cols=4)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    headers = ["#", "Outil", "Rôle principal", "Install."]
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ""
        r = cell.paragraphs[0].add_run(h)
        set_run_font(r, size=9, bold=True, color=WHITE)
        shade_cell(cell, "0D3B66")
        set_cell_border(cell)
    overview = [
        ("01", "OSINT Framework", "Carte des ressources", "Navigateur"),
        ("02", "SpiderFoot", "Automatisation multi-sources", "Python"),
        ("03", "theHarvester", "E-mails / hosts / sous-domaines", "CLI / Kali"),
        ("04", "Recon-ng", "Framework d’enquête modulaire", "Python"),
        ("05", "OWASP Amass", "Surface d’attaque DNS", "Go / Kali"),
        ("06", "Sherlock", "Pseudos / SOCMINT", "Python"),
        ("07", "ExifTool", "Métadonnées fichiers", "Paquet OS"),
        ("08", "Maltego CE", "Graphe relationnel", "Client + compte"),
        ("09", "crt.sh", "Certificats / sous-domaines", "Web / API"),
        ("10", "Wayback Machine", "Archives web historiques", "Web + CLI"),
    ]
    for i, row in enumerate(overview):
        cells = table.add_row().cells
        bg = "E8F0F8" if i % 2 == 0 else "FFFFFF"
        for j, val in enumerate(row):
            cells[j].text = ""
            r = cells[j].paragraphs[0].add_run(val)
            set_run_font(r, size=9, bold=(j == 1), color=DARK)
            shade_cell(cells[j], bg)
            set_cell_border(cells[j], color="BDD3E6")
    doc.add_paragraph()

    # ========== 3 FICHES ==========
    add_heading_styled(doc, "3. Fiches détaillées", 1)
    add_para(
        doc,
        "Chaque fiche suit le même plan : rôle, fiche d’identité, installation, usage, "
        "exemples, cas d’emploi cyber, limites.",
        align="justify",
        italic=True,
        color=GRAY,
    )

    for t in TOOLS:
        add_heading_styled(doc, f"3.{int(t['num'])}  {t['name']}", 2)

        add_kv_table(
            doc,
            [
                ("Type", t["type"]),
                ("Licence / modèle", t["licence"]),
                ("URL officielle", t["url"]),
            ],
            headers=("Fiche d’identité", "Valeur"),
        )

        add_heading_styled(doc, "Rôle pour l’expert", 3)
        add_para(doc, t["role"], align="justify")

        add_heading_styled(doc, "Installation", 3)
        for line in t["install"]:
            add_bullet(doc, line)
        add_para(doc, "Commandes :", size=10, bold=True, color=NAVY, space_after=2)
        add_code(doc, t["install_cmd"])

        add_heading_styled(doc, "Usage opérationnel", 3)
        for i, u in enumerate(t["usage"], 1):
            add_bullet(doc, u, bold_prefix=f"{i}. ")

        if t.get("exemples"):
            add_para(doc, "Exemples :", size=10, bold=True, color=NAVY, space_after=2)
            add_code(doc, t["exemples"])

        add_heading_styled(doc, "Cas d’emploi cybersécurité", 3)
        for c in t["cas"]:
            add_bullet(doc, c)

        add_heading_styled(doc, "Limites et précautions", 3)
        add_para(doc, t["limites"], align="justify")

    # ========== 4 STACK ==========
    add_heading_styled(doc, "4. Stack minimale recommandée", 1)
    add_para(
        doc,
        "Pour un expert (SOC, CERT, pentester, analyste CTI), l’enchaînement suivant est efficace :",
        align="justify",
    )
    steps = [
        ("Cadrer", "OSINT Framework — choisir les branches pertinentes."),
        ("Découvrir l’infra", "crt.sh + Amass (passif) + theHarvester."),
        ("Historique", "Wayback Machine / waybackurls."),
        ("Automatiser", "SpiderFoot (scan passif) et/ou Recon-ng (workspace)."),
        ("Identité", "Sherlock (pseudos) + ExifTool (fichiers/images)."),
        ("Synthétiser", "Maltego CE — graphe et livrable visuel."),
    ]
    table = doc.add_table(rows=1, cols=2)
    hdr = table.rows[0].cells
    for i, h in enumerate(("Étape", "Outils")):
        hdr[i].text = ""
        r = hdr[i].paragraphs[0].add_run(h)
        set_run_font(r, size=10, bold=True, color=WHITE)
        shade_cell(hdr[i], "1F4E79")
        set_cell_border(hdr[i])
    for i, (a, b) in enumerate(steps):
        row = table.add_row().cells
        bg = "F7FAFC" if i % 2 == 0 else "FFFFFF"
        row[0].text = ""
        r = row[0].paragraphs[0].add_run(a)
        set_run_font(r, size=10, bold=True, color=NAVY)
        row[1].text = ""
        r = row[1].paragraphs[0].add_run(b)
        set_run_font(r, size=10, color=DARK)
        shade_cell(row[0], bg)
        shade_cell(row[1], bg)
        set_cell_border(row[0], color="BDD3E6")
        set_cell_border(row[1], color="BDD3E6")
    doc.add_paragraph()

    add_para(doc, "Environnement de travail type :", bold=True, color=NAVY, space_after=4)
    add_code(
        doc,
        "# Machine d’enquête dédiée (Kali ou VM Linux isolée)\n"
        "sudo apt update\n"
        "sudo apt install -y theharvester amass maltego git python3-pip pipx \\\n"
        "  libimage-exiftool-perl\n"
        "pipx install sherlock-project\n"
        "git clone https://github.com/smicallef/spiderfoot.git\n"
        "git clone https://github.com/lanmaster53/recon-ng.git",
    )

    # ========== 5 COMPLEMENTS ==========
    add_heading_styled(doc, "5. Outils complémentaires (à connaître)", 1)
    add_para(
        doc,
        "Hors du « top 10 » strictement open source / public, un expert croise souvent :",
        align="justify",
    )
    comps = [
        ("Shodan", "Moteur de services Internet exposés (freemium). https://www.shodan.io"),
        ("Censys", "Assets Internet et certificats. https://search.censys.io"),
        ("Have I Been Pwned", "Vérification de fuites d’identifiants. https://haveibeenpwned.com"),
        ("VirusTotal", "Réputation fichiers / URL / domaines. https://www.virustotal.com"),
        ("Holehe", "E-mail → comptes associés (OSS). https://github.com/megadose/holehe"),
        ("urlscan.io", "Analyse de pages / phishing. https://urlscan.io"),
    ]
    for name, desc in comps:
        add_bullet(doc, f" {desc}", bold_prefix=name)

    # ========== 6 BONNES PRATIQUES ==========
    add_heading_styled(doc, "6. Bonnes pratiques d’expert", 1)
    add_bullet(doc, "Séparer le navigateur d’enquête (profil dédié, pas de compte perso).")
    add_bullet(doc, "Journaliser : date, outil, requête, hash des exports.")
    add_bullet(doc, "Vérifier en multi-sources (un seul outil ne « prouve » rien).")
    add_bullet(doc, "Préférer le passif ; l’actif uniquement si le contrat le permet.")
    add_bullet(doc, "Gérer les clés API dans un coffre, jamais dans Git.")
    add_bullet(doc, "Anonymiser / minimiser les extraits dans les rapports diffusés.")
    add_bullet(doc, "Mettre à jour les outils (Amass, Sherlock, SpiderFoot évoluent vite).")

    add_callout(
        doc,
        "Éthique.",
        "L’OSINT n’est pas de la surveillance illégale ni du doxxing. "
        "L’expert documente la finalité, respecte les personnes et la loi applicable "
        "(code pénal, RGPD, mandat d’audit).",
    )

    # ========== 7 REFS ==========
    add_heading_styled(doc, "7. Références", 1)
    refs = [
        "OSINT Framework — https://osintframework.com",
        "SpiderFoot — https://github.com/smicallef/spiderfoot",
        "theHarvester — https://github.com/laramies/theHarvester",
        "Recon-ng — https://github.com/lanmaster53/recon-ng",
        "OWASP Amass — https://github.com/owasp-amass/amass",
        "Sherlock — https://github.com/sherlock-project/sherlock",
        "ExifTool — https://exiftool.org",
        "Maltego — https://www.maltego.com",
        "crt.sh — https://crt.sh",
        "Internet Archive Wayback Machine — https://web.archive.org",
        "OWASP — Attack Surface Management / reconnaissance (documentation projet Amass)",
    ]
    for ref in refs:
        add_bullet(doc, ref)

    add_para(doc, "")
    line = doc.add_table(rows=1, cols=1)
    cell = line.cell(0, 0)
    shade_cell(cell, "0D3B66")
    set_cell_border(cell, color="0D3B66")
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Fin du document  •  ABOU SAYABOU  •  Guide OSINT v1.0")
    set_run_font(r, size=10, bold=True, color=WHITE)

    add_para(
        doc,
        "\nLes estimations, captures et commandes sont fournies à titre pédagogique. "
        "L’auteur décline toute responsabilité en cas d’usage hors cadre légal ou contractuel.",
        size=9,
        italic=True,
        color=GRAY,
        align="center",
    )

    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT)
    print(f"Document généré : {OUT}")


if __name__ == "__main__":
    build()
