#!/usr/bin/env python3
"""Génère le pack Word + Excel NigerCertify (plan com & calendrier éditorial)."""

from datetime import date, timedelta
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor
from openpyxl import Workbook
from openpyxl.styles import Alignment, Border, Font, PatternFill, Side
from openpyxl.utils import get_column_letter
from openpyxl.worksheet.datavalidation import DataValidation

OUT = Path(__file__).resolve().parent
WHATSAPP = "94 10 70 74"
WA_LINK = "https://wa.me/22794107074"


def add_heading(doc: Document, text: str, level: int = 1) -> None:
    doc.add_heading(text, level=level)


def add_para(doc: Document, text: str, bold: bool = False) -> None:
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(11)


def add_bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        doc.add_paragraph(item, style="List Bullet")


def build_word() -> Path:
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = title.add_run("NigerCertify")
    r.bold = True
    r.font.size = Pt(28)
    r.font.color.rgb = RGBColor(0x0F, 0x4C, 0x3A)

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    s = sub.add_run("Plan de communication & catalogue formations\nVersion 2.3 — Niger & Sénégal")
    s.font.size = Pt(14)

    add_para(doc, f"WhatsApp commercial : {WHATSAPP}  ·  {WA_LINK}", bold=True)
    add_para(doc, "Budget com : 20 000 FCFA / mois — stratégie organique (Facebook, LinkedIn, YouTube, WhatsApp).")

    add_heading(doc, "1. Positionnement", 1)
    add_para(
        doc,
        "NigerCertify forme et certifie (PECB + LPI) et digitalise les métiers via nBusiness et CouturePro.",
    )
    add_para(doc, "Promesse : Certifiez-vous. Structurez votre métier.", bold=True)
    add_bullets(
        doc,
        [
            "PECB : ISO/IEC 27001, ISO/IEC 27701 — session en ligne 10 octobre 2026 — 350 000 FCFA (accès 12 mois, 2 retakes)",
            "LPI : Linux Essentials 150 000 FCFA · LPIC-1/2/3 200 000 FCFA / niveau — format mixte",
            "nBusiness : gestion multi-secteur (commerce, éducation, business, pressing…)",
            "CouturePro : gestion ateliers de couture",
            "Lab GitHub offensif : hors communication grand public",
        ],
    )

    add_heading(doc, "2. Architecture d’offres (répartition com)", 1)
    table = doc.add_table(rows=5, cols=3)
    table.style = "Table Grid"
    headers = ("Priorité", "Offre", "Part des messages")
    for i, h in enumerate(headers):
        table.rows[0].cells[i].text = h
    data = [
        ("#1", "PECB (ISO 27001 / 27701)", "60 %"),
        ("#2", "LPI LE → LPIC-3", "20 %"),
        ("#3", "nBusiness", "15 %"),
        ("#4", "CouturePro", "5 %"),
    ]
    for r_i, row in enumerate(data, start=1):
        for c_i, val in enumerate(row):
            table.rows[r_i].cells[c_i].text = val

    add_heading(doc, "3. Catalogue formations", 1)
    add_heading(doc, "3.1 PECB", 2)
    add_bullets(
        doc,
        [
            "Normes : ISO/IEC 27001 · ISO/IEC 27701",
            "Démarrage : 10 octobre 2026",
            "Modalité : 100 % en ligne",
            "Accès plateforme : 12 mois",
            "Examens : inclus + 2 retakes",
            "Tarif : 350 000 FCFA",
            "Inscription : WhatsApp 94 10 70 74 — JE M’INSCRIS",
        ],
    )

    add_heading(doc, "3.2 LPI (Linux)", 2)
    t2 = doc.add_table(rows=5, cols=4)
    t2.style = "Table Grid"
    for i, h in enumerate(("Niveau", "Certification", "Tarif", "Format")):
        t2.rows[0].cells[i].text = h
    lpi_rows = [
        ("LPI LE", "Linux Essentials", "150 000 FCFA", "Mixte"),
        ("LPIC-1", "Linux Administrator", "200 000 FCFA", "Mixte"),
        ("LPIC-2", "Linux Engineer", "200 000 FCFA", "Mixte"),
        ("LPIC-3", "Linux Enterprise", "200 000 FCFA", "Mixte"),
    ]
    for r_i, row in enumerate(lpi_rows, start=1):
        for c_i, val in enumerate(row):
            t2.rows[r_i].cells[c_i].text = val
    add_para(doc, "Mot-clé WhatsApp : LPI + niveau (LE / LPIC1 / LPIC2 / LPIC3)")

    add_heading(doc, "4. Canaux prioritaires", 1)
    add_bullets(
        doc,
        [
            "WhatsApp Business — canal #1 commercial",
            "LinkedIn — B2B / PECB / LPI",
            "Facebook — PME, CouturePro, audiences locales Niger & Sénégal",
            "YouTube — démos courtes (PECB tips, Linux, nBusiness, CouturePro)",
            "Email — nurturing sessions",
        ],
    )

    add_heading(doc, "5. Mots-clés WhatsApp", 1)
    t3 = doc.add_table(rows=7, cols=2)
    t3.style = "Table Grid"
    t3.rows[0].cells[0].text = "Mot-clé"
    t3.rows[0].cells[1].text = "Action"
    keywords = [
        ("PROGRAMME", "Catalogue complet"),
        ("PECB / JE M’INSCRIS", "Session ISO 10/10/2026"),
        ("LPI", "Parcours Linux"),
        ("LE / LPIC1 / LPIC2 / LPIC3", "Niveau LPI précis"),
        ("NBUSINESS", "Démo nBusiness"),
        ("COUTUREPRO", "Démo CouturePro"),
    ]
    for r_i, (k, v) in enumerate(keywords, start=1):
        t3.rows[r_i].cells[0].text = k
        t3.rows[r_i].cells[1].text = v

    add_heading(doc, "6. Scripts WhatsApp (extraits)", 1)
    add_para(doc, "Accueil général", bold=True)
    add_para(
        doc,
        "Bonjour, NigerCertify — formations certifiantes (Niger & Sénégal).\n"
        "1) PECB ISO 27001/27701 — 10 oct. 2026 — 350 000 FCFA — 12 mois — 2 retakes\n"
        "2) LPI mixte — LE 150 000 FCFA — LPIC-1/2/3 200 000 FCFA / niveau\n"
        "Aussi : nBusiness · CouturePro — WhatsApp 94 10 70 74",
    )

    add_heading(doc, "7. Règles de publication", 1)
    add_bullets(
        doc,
        [
            "1 contenu = 1 public + 1 offre + 1 CTA",
            "Chaque publication = Titre + Contenu + Hashtags (voir Excel)",
            "Répartition : PECB 60 % · LPI 20 % · nBusiness 15 % · CouturePro 5 %",
            "Pas de contenu lab offensif / webshell sur Facebook, LinkedIn, YouTube grand public",
            "Toujours indiquer WhatsApp 94 10 70 74 et un mot-clé clair",
            "Calendrier détaillé : fichier Excel joint (Facebook / LinkedIn / YouTube)",
        ],
    )

    add_heading(doc, "8. Feuille de route 90 jours", 1)
    add_bullets(
        doc,
        [
            "Jours 1–30 : catalogue PDF, bios, 9 posts LinkedIn, liste 50+50 prospects, WhatsApp structuré",
            "Jours 31–60 : pousser session PECB 10/10, 2 témoignages, démos LPI/nBusiness",
            "Jours 61–90 : calendrier T2, partenariat, industrialiser les relances",
        ],
    )

    add_heading(doc, "9. Contact", 1)
    add_para(doc, f"WhatsApp : {WHATSAPP}", bold=True)
    add_para(doc, WA_LINK)
    add_para(doc, "Marchés : Niger · Sénégal · distanciel francophone")

    footer = doc.add_paragraph()
    footer.add_run(
        "\nDocument généré pour NigerCertify — usage interne / commercial. "
        "Compléter modalités de paiement et calendrier de démarrage LPI."
    ).italic = True

    path = OUT / "NigerCertify-Plan-Communication.docx"
    doc.save(path)
    return path


# --- Excel calendar ---
# Chaque post = (titre, contenu, hashtags, cta)

THEMES = {
    "PECB": {
        "fill": "1F6B4F",
        "posts": [
            (
                "Session PECB — inscriptions ouvertes (10 oct. 2026)",
                "Les inscriptions sont ouvertes pour la session PECB NigerCertify.\n\n"
                "📅 Démarrage : 10 octobre 2026\n"
                "💻 100 % en ligne\n"
                "📘 ISO/IEC 27001 · ISO/IEC 27701\n"
                "⏱ Accès plateforme 12 mois\n"
                "🔁 2 retakes examens inclus\n"
                "💰 Tarif : 350 000 FCFA\n\n"
                "Niger & Sénégal — certification internationale, accompagnement local.\n\n"
                f"👉 WhatsApp {WHATSAPP} — écrivez JE M’INSCRIS",
                "#PECB #ISO27001 #ISO27701 #Cybersécurité #Certification #NigerCertify #Niger #Sénégal #FormationEnLigne",
                "JE M’INSCRIS",
            ),
            (
                "Pourquoi se certifier ISO 27001 ?",
                "ISO 27001, c’est le langage commun de la sécurité de l’information.\n\n"
                "Avec NigerCertify, vous préparez votre certification PECB :\n"
                "• Session en ligne dès le 10 octobre 2026\n"
                "• Accès 12 mois + 2 retakes\n"
                "• Tarif : 350 000 FCFA\n\n"
                "Idéal pour RSSI, IT, consultants et équipes conformité.\n\n"
                f"📩 Demandez le programme au {WHATSAPP} (mot-clé PROGRAMME)",
                "#ISO27001 #PECB #SécuritéDeLInformation #SMSI #FormationIT #NigerCertify #AfriqueDeLOuest",
                "PROGRAMME",
            ),
            (
                "350 000 FCFA : ce qui est vraiment inclus",
                "Pour 350 000 FCFA, la session PECB du 10/10/2026 inclut :\n\n"
                "✅ Formation PECB (ISO 27001 / 27701)\n"
                "✅ Accès en ligne 12 mois\n"
                "✅ Passage d’examen + 2 retakes\n"
                "✅ Accompagnement NigerCertify\n\n"
                "Pas de surprise. Un parcours clair pour réussir.\n\n"
                f"WhatsApp {WHATSAPP} — répondez PECB",
                "#PECB #Formation #ISO27001 #ISO27701 #CertificationPro #NigerCertify #InvestissementFormation",
                "PECB",
            ),
            (
                "ISO 27701 : la privacy au cœur du SMSI",
                "Protéger les données personnelles n’est plus optionnel.\n\n"
                "ISO/IEC 27701 complète ISO 27001 pour structurer la privacy.\n"
                "Session PECB NigerCertify — démarrage 10 octobre 2026\n"
                "En ligne · 12 mois · 2 retakes · 350 000 FCFA\n\n"
                f"Places limitées → {WHATSAPP}",
                "#ISO27701 #Privacy #ProtectionDesDonnées #RGPD #PECB #NigerCertify #Conformité",
                "PROGRAMME",
            ),
            (
                "Niamey ↔ Dakar : une même session en ligne",
                "Que vous soyez au Niger ou au Sénégal, la session PECB est la même :\n\n"
                "• Exigence internationale\n"
                "• Accompagnement local NigerCertify\n"
                "• Démarrage 10 octobre 2026\n"
                "• 350 000 FCFA · 2 retakes\n\n"
                f"Inscrivez-vous : WhatsApp {WHATSAPP}",
                "#Niger #Sénégal #PECB #FormationEnLigne #ISO27001 #NigerCertify #Certification",
                "JE M’INSCRIS",
            ),
        ],
    },
    "LPI": {
        "fill": "2E5A88",
        "posts": [
            (
                "Parcours LPI : de Linux Essentials à LPIC-3",
                "Admin Linux : un chemin clair avec NigerCertify.\n\n"
                "🐧 LPI LE (Linux Essentials) — 150 000 FCFA\n"
                "🐧 LPIC-1 · LPIC-2 · LPIC-3 — 200 000 FCFA / niveau\n"
                "📍 Format mixte (présentiel + en ligne)\n\n"
                "Progression : LE → LPIC-1 → LPIC-2 → LPIC-3\n\n"
                f"Répondez LPI au {WHATSAPP} + votre niveau",
                "#LPI #LPIC #Linux #LinuxEssentials #SysAdmin #NigerCertify #FormationLinux #OpenSource",
                "LPI",
            ),
            (
                "Linux Essentials — le bon premier pas",
                "Vous débutez sous Linux ou vous vous reconvertissez en IT ?\n\n"
                "LPI Linux Essentials (LPI LE) chez NigerCertify :\n"
                "• Bases système & ligne de commande\n"
                "• Culture open source\n"
                "• Tarif : 150 000 FCFA\n"
                "• Format mixte\n\n"
                f"Message LPI LE au {WHATSAPP}",
                "#LinuxEssentials #LPILE #Linux #ReconversionIT #Formation #NigerCertify #LPIC",
                "LPI LE",
            ),
            (
                "LPIC-1, LPIC-2, LPIC-3 : montez en niveau",
                "Déjà à l’aise sous Linux ? Passez au niveau supérieur.\n\n"
                "• LPIC-1 — Administrateur\n"
                "• LPIC-2 — Engineer\n"
                "• LPIC-3 — Enterprise\n\n"
                "Tarif : 200 000 FCFA / niveau · Format mixte\n"
                "NigerCertify — Niger & Sénégal\n\n"
                f"WhatsApp {WHATSAPP} — mot-clé LPI",
                "#LPIC1 #LPIC2 #LPIC3 #LinuxAdmin #DevOps #NigerCertify #CertificationLinux",
                "LPI",
            ),
            (
                "Tip Linux : une commande, un métier",
                "Maîtriser Linux, c’est maîtriser le terrain IT.\n\n"
                "Chaque semaine, NigerCertify vous prépare aux certifications LPI :\n"
                "LE (150 000 FCFA) · LPIC-1/2/3 (200 000 FCFA) · format mixte.\n\n"
                "Aujourd’hui : prenez 10 minutes pour réviser une commande.\n"
                "Demain : visez la certif.\n\n"
                f"Écrivez LPI au {WHATSAPP}",
                "#LinuxTip #LPI #SysAdmin #Terminal #OpenSource #NigerCertify #FormationIT",
                "LPI",
            ),
        ],
    },
    "nBusiness": {
        "fill": "8A6D1D",
        "posts": [
            (
                "nBusiness — la gestion multi-secteur",
                "Commerce, école, pressing, business…\n\n"
                "nBusiness centralise :\n"
                "• Stocks\n"
                "• Clients\n"
                "• Ventes\n\n"
                "Sans usine à gaz. Pensé pour les PME du Niger et du Sénégal.\n\n"
                f"Démo 10 min : écrivez NBUSINESS au {WHATSAPP}",
                "#nBusiness #GestionPME #Digitalisation #Commerce #Pressing #Éducation #NigerCertify #PME",
                "NBUSINESS",
            ),
            (
                "Démo nBusiness en 10 minutes",
                "Envie de voir nBusiness en action ?\n\n"
                "On vous montre en 10 minutes sur WhatsApp comment suivre stock, clients et ventes.\n"
                "Idéal commerce, école, pressing et business.\n\n"
                f"Réservez : {WHATSAPP} — mot-clé NBUSINESS",
                "#nBusiness #DémoLogiciel #PME #GestionCommerciale #Niger #Sénégal #NigerCertify",
                "NBUSINESS",
            ),
        ],
    },
    "CouturePro": {
        "fill": "8B3A4A",
        "posts": [
            (
                "CouturePro — le logiciel des ateliers de couture",
                "Les ateliers perdent du temps sans suivi clair.\n\n"
                "CouturePro vous aide à gérer :\n"
                "• Commandes clients\n"
                "• Stocks tissus & fournitures\n"
                "• Délais de livraison\n\n"
                "Un outil simple pour les couturiers et stylistes.\n\n"
                f"Démo : écrivez COUTUREPRO au {WHATSAPP}",
                "#CouturePro #AtelierCouture #Mode #Artisanat #Digitalisation #NigerCertify #GestionAtelier",
                "COUTUREPRO",
            ),
        ],
    },
}


def next_weekday(d: date, weekday: int) -> date:
    """weekday: 0=Mon ... 6=Sun"""
    days = (weekday - d.weekday()) % 7
    return d + timedelta(days=days)


def pick_post(offre: str, counters: dict[str, int]) -> tuple[str, str, str, str]:
    posts = THEMES[offre]["posts"]
    i = counters[offre] % len(posts)
    counters[offre] += 1
    return posts[i]


def build_excel(start: date | None = None) -> Path:
    """Calendrier 12 semaines — chaque ligne = Titre + Contenu + Hashtags."""
    if start is None:
        today = date.today()
        start = next_weekday(today, 0)
        if start == today and today.weekday() != 0:
            start = next_weekday(today + timedelta(days=1), 0)

    wb = Workbook()

    header_fill = PatternFill("solid", fgColor="0F4C3A")
    header_font = Font(bold=True, color="FFFFFF", name="Calibri", size=11)
    thin = Border(
        left=Side(style="thin", color="CCCCCC"),
        right=Side(style="thin", color="CCCCCC"),
        top=Side(style="thin", color="CCCCCC"),
        bottom=Side(style="thin", color="CCCCCC"),
    )
    wrap = Alignment(wrap_text=True, vertical="top")

    guide = wb.active
    guide.title = "Guide"
    guide["A1"] = "Calendrier de publication NigerCertify"
    guide["A1"].font = Font(bold=True, size=16, color="0F4C3A")
    guide["A2"] = f"WhatsApp : {WHATSAPP}  |  {WA_LINK}"
    guide["A3"] = "Canaux : Facebook · LinkedIn · YouTube"
    guide["A4"] = "Chaque publication contient : TITRE + CONTENU + HASHTAGS"
    guide["A5"] = "Répartition : PECB 60% · LPI 20% · nBusiness 15% · CouturePro 5%"
    guide["A7"] = "Colonnes clés"
    for i, line in enumerate(
        [
            "Titre = accroche à afficher / titre YouTube",
            "Contenu = texte prêt à copier-coller (ou script YouTube)",
            "Hashtags = à coller en fin de post (Facebook/LinkedIn) ou dans la description YouTube",
            "CTA = mot-clé WhatsApp",
            "Statut = À faire / En cours / Publié / Reporté",
            "Ne jamais publier de contenu lab offensif sur ces canaux.",
        ],
        start=8,
    ):
        guide[f"A{i}"] = line
    guide.column_dimensions["A"].width = 110

    ws = wb.create_sheet("Calendrier", 0)
    headers = [
        "Semaine",
        "Date",
        "Jour",
        "Canal",
        "Offre",
        "Titre",
        "Contenu",
        "Hashtags",
        "Format",
        "CTA / Mot-clé",
        "Statut",
        "Responsable",
        "Lien publication",
        "Notes",
    ]
    for col, h in enumerate(headers, 1):
        cell = ws.cell(1, col, h)
        cell.fill = header_fill
        cell.font = header_font
        cell.alignment = Alignment(wrap_text=True, vertical="center", horizontal="center")
        cell.border = thin

    week_slots = [
        (0, "LinkedIn", "PECB", "Post texte / carrousel"),
        (1, "Facebook", "LPI", "Post image + texte"),
        (2, "LinkedIn", "PECB", "Post texte + CTA"),
        (3, "YouTube", "PECB", "Vidéo 5–8 min ou Short"),
        (4, "Facebook", "nBusiness", "Post / Reels"),
        (6, "Facebook", "PECB", "Post CTA session"),
    ]

    counters = {"PECB": 0, "LPI": 0, "nBusiness": 0, "CouturePro": 0}
    row = 2
    for week in range(1, 13):
        week_start = start + timedelta(weeks=week - 1)
        for day_offset, canal, offre_default, fmt in week_slots:
            pub_date = week_start + timedelta(days=day_offset)

            offre = offre_default
            if canal == "YouTube" and week % 2 == 0:
                offre = "LPI"
            if canal == "Facebook" and day_offset == 4 and week % 4 == 0:
                offre = "CouturePro"

            title, contenu, hashtags, cta = pick_post(offre, counters)
            bundle = THEMES[offre]

            if canal == "YouTube":
                contenu = (
                    f"TITRE VIDÉO : {title}\n\n"
                    f"HOOK (0–5 s) : {title}\n\n"
                    f"SCRIPT / CONTENU :\n{contenu}\n\n"
                    f"CTA ÉCRAN DE FIN : Écrivez {cta} au {WHATSAPP}\n\n"
                    f"DESCRIPTION YOUTUBE :\n{contenu.split(chr(10))[0]}\n"
                    f"WhatsApp : {WA_LINK}?text={cta.replace(' ', '%20')}\n"
                    f"NigerCertify · Niger & Sénégal"
                )
                fmt = "YouTube Short (45–60s)" if week % 3 == 0 else "Vidéo 5–8 min"

            jour = ["Lun", "Mar", "Mer", "Jeu", "Ven", "Sam", "Dim"][pub_date.weekday()]
            values = [
                week,
                pub_date.isoformat(),
                jour,
                canal,
                offre,
                title,
                contenu,
                hashtags,
                fmt,
                cta,
                "À faire",
                "",
                "",
                "",
            ]
            fill = PatternFill("solid", fgColor=bundle["fill"])
            light = PatternFill(
                "solid",
                fgColor=(
                    "E8F5F0"
                    if offre == "PECB"
                    else "EEF3F8"
                    if offre == "LPI"
                    else "F8F3E0"
                    if offre == "nBusiness"
                    else "F8E8EC"
                ),
            )

            for col, val in enumerate(values, 1):
                cell = ws.cell(row, col, val)
                cell.border = thin
                cell.alignment = wrap
                if col == 5:
                    cell.fill = fill
                    cell.font = Font(bold=True, color="FFFFFF")
                elif col == 4:
                    cell.fill = light
                elif col in (6, 7, 8):  # Titre, Contenu, Hashtags
                    cell.font = Font(name="Calibri", size=10)
            ws.row_dimensions[row].height = 90
            row += 1

    widths = {
        "A": 9,
        "B": 12,
        "C": 6,
        "D": 11,
        "E": 11,
        "F": 38,
        "G": 58,
        "H": 42,
        "I": 22,
        "J": 14,
        "K": 11,
        "L": 12,
        "M": 24,
        "N": 16,
    }
    for col, w in widths.items():
        ws.column_dimensions[col].width = w
    ws.row_dimensions[1].height = 32
    ws.freeze_panes = "A2"
    ws.auto_filter.ref = f"A1:N{row - 1}"

    dv = DataValidation(type="list", formula1='"À faire,En cours,Publié,Reporté"', allow_blank=True)
    ws.add_data_validation(dv)
    dv.add(f"K2:K{row - 1}")

    dv_canal = DataValidation(type="list", formula1='"Facebook,LinkedIn,YouTube"', allow_blank=True)
    ws.add_data_validation(dv_canal)
    dv_canal.add(f"D2:D{row - 1}")

    dv_offre = DataValidation(type="list", formula1='"PECB,LPI,nBusiness,CouturePro"', allow_blank=True)
    ws.add_data_validation(dv_offre)
    dv_offre.add(f"E2:E{row - 1}")

    recap = wb.create_sheet("Recap_mensuel")
    recap["A1"] = "Objectif mensuel (indicatif)"
    recap["A1"].font = Font(bold=True, size=14, color="0F4C3A")
    for col, h in enumerate(("Canal", "Posts / mois (cible)", "Notes"), 1):
        c = recap.cell(3, col, h)
        c.fill = header_fill
        c.font = header_font
    for r_i, row_data in enumerate(
        [
            ("LinkedIn", "8–10", "Chaque post = Titre + Contenu + Hashtags"),
            ("Facebook", "10–12", "Coller Hashtags en fin de publication"),
            ("YouTube", "4–6", "Titre = titre vidéo · Contenu = script · Hashtags = description"),
            ("WhatsApp Status", "20+", "Hors Excel mais quotidien"),
        ],
        start=4,
    ):
        for c_i, val in enumerate(row_data, 1):
            recap.cell(r_i, c_i, val).border = thin
    recap.column_dimensions["A"].width = 18
    recap.column_dimensions["B"].width = 22
    recap.column_dimensions["C"].width = 55

    offres = wb.create_sheet("Tarifs_offres")
    offres["A1"] = "Grille tarifaire formations"
    offres["A1"].font = Font(bold=True, size=14, color="0F4C3A")
    for col, h in enumerate(("Offre", "Détail", "Format", "Tarif (FCFA)", "CTA"), 1):
        c = offres.cell(3, col, h)
        c.fill = header_fill
        c.font = header_font
    for r_i, row_data in enumerate(
        [
            ("PECB ISO 27001/27701", "Session 10 oct. 2026 · 12 mois · 2 retakes", "En ligne", 350000, "JE M’INSCRIS"),
            ("LPI LE", "Linux Essentials", "Mixte", 150000, "LPI LE"),
            ("LPIC-1", "Linux Administrator", "Mixte", 200000, "LPIC1"),
            ("LPIC-2", "Linux Engineer", "Mixte", 200000, "LPIC2"),
            ("LPIC-3", "Linux Enterprise", "Mixte", 200000, "LPIC3"),
            ("nBusiness", "Gestion multi-secteur", "Démo", "Sur devis", "NBUSINESS"),
            ("CouturePro", "Ateliers couture", "Démo", "Sur devis", "COUTUREPRO"),
        ],
        start=4,
    ):
        for c_i, val in enumerate(row_data, 1):
            cell = offres.cell(r_i, c_i, val)
            cell.border = thin
            if c_i == 4 and isinstance(val, int):
                cell.number_format = "#,##0"
    for col in range(1, 6):
        offres.column_dimensions[get_column_letter(col)].width = 28

    path = OUT / "NigerCertify-Calendrier-Publication.xlsx"
    wb.save(path)
    return path


def main() -> None:
    word = build_word()
    excel = build_excel()
    print(f"Word  : {word}")
    print(f"Excel : {excel}")


if __name__ == "__main__":
    main()
