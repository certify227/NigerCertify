#!/usr/bin/env python3
"""Convertit un fichier Markdown en document Word (.docx)."""

from pathlib import Path
from docx import Document


def add_markdown_table(doc: Document, table_lines: list[str]) -> None:
    rows = []
    for line in table_lines:
        stripped = line.strip()
        if not stripped.startswith("|"):
            continue
        parts = [p.strip() for p in stripped.strip("|").split("|")]
        rows.append(parts)

    if len(rows) < 2:
        return

    # Ignore separator row (----)
    header = rows[0]
    data_rows = rows[2:] if len(rows) >= 2 else []
    cols = len(header)
    table = doc.add_table(rows=1, cols=cols)
    table.style = "Table Grid"
    hdr_cells = table.rows[0].cells
    for i, value in enumerate(header):
        hdr_cells[i].text = value

    for row in data_rows:
        tr = table.add_row().cells
        for i in range(cols):
            tr[i].text = row[i] if i < len(row) else ""


def md_to_docx(md_path: Path, docx_path: Path) -> None:
    doc = Document()
    lines = md_path.read_text(encoding="utf-8").splitlines()

    i = 0
    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        if not stripped or stripped == "---":
            doc.add_paragraph("")
            i += 1
            continue

        if stripped.startswith("|"):
            table_lines = []
            while i < len(lines) and lines[i].strip().startswith("|"):
                table_lines.append(lines[i])
                i += 1
            add_markdown_table(doc, table_lines)
            continue

        if stripped.startswith("### "):
            doc.add_heading(stripped[4:], level=3)
            i += 1
            continue
        if stripped.startswith("## "):
            doc.add_heading(stripped[3:], level=2)
            i += 1
            continue
        if stripped.startswith("# "):
            doc.add_heading(stripped[2:], level=1)
            i += 1
            continue

        if stripped.startswith("- "):
            doc.add_paragraph(stripped[2:], style="List Bullet")
            i += 1
            continue

        if (
            len(stripped) > 2
            and stripped[0].isdigit()
            and stripped.split(".", 1)[0].isdigit()
            and ". " in stripped
        ):
            content = stripped.split(". ", 1)[1]
            doc.add_paragraph(content, style="List Number")
            i += 1
            continue

        # Nettoyage léger markdown inline
        text = stripped.replace("**", "").replace("`", "")
        doc.add_paragraph(text)
        i += 1

    doc.save(docx_path)


def main() -> None:
    base = Path(__file__).resolve().parent
    md_file = base / "Analyse_Technique_SIEM_3_Offres.md"
    docx_file = base / "Analyse_Technique_SIEM_3_Offres.docx"
    md_to_docx(md_file, docx_file)
    print(f"Document Word généré : {docx_file}")


if __name__ == "__main__":
    main()
